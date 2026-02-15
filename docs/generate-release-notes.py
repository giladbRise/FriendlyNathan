#!/usr/bin/env python3
"""Generate Friendly Nathan v1.0 Release Notes as a styled DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color palette (matches Friendly Nathan theme) ──────────────────
CORAL = RGBColor(0xF0, 0x6A, 0x3E)       # Primary coral-orange
TEAL = RGBColor(0x14, 0xB8, 0xA6)        # Secondary teal
GOLD = RGBColor(0xE8, 0xAA, 0x0A)        # Accent yellow
DARK = RGBColor(0x2D, 0x1F, 0x14)        # Dark brown text
MEDIUM = RGBColor(0x5A, 0x4D, 0x44)      # Medium brown
LIGHT_TEXT = RGBColor(0x8A, 0x74, 0x68)   # Muted text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xFB, 0xF7, 0xF0)    # Warm cream
CORAL_LIGHT = RGBColor(0xFF, 0xF0, 0xEB)  # Light coral bg
TEAL_LIGHT = RGBColor(0xE6, 0xFB, 0xF8)  # Light teal bg
GOLD_LIGHT = RGBColor(0xFF, 0xF8, 0xE1)  # Light gold bg

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1, color=CORAL, space_before=18, space_after=6):
    """Add a heading with custom styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    return h

def add_body(doc, text, bold=False, color=DARK, size=10.5, space_after=6, alignment=None):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p

def add_bullet(doc, text, level=0, bold_prefix=None, color=DARK):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = color
        run.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = MEDIUM
        run.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = MEDIUM
        run.font.name = 'Calibri'
    return p

def add_feature_table(doc, rows, headers=None, col_widths=None):
    """Add a styled table."""
    num_cols = len(rows[0]) if rows else len(headers)
    has_header = headers is not None
    total_rows = len(rows) + (1 if has_header else 0)
    table = doc.add_table(rows=total_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    if has_header:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = WHITE
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'F06A3E')

    # Data rows
    start = 1 if has_header else 0
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[start + r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            run.font.name = 'Calibri'
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'FFF5F0')

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # Spacer
    return table

def add_divider(doc):
    """Add a visual divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('━' * 72)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xE8, 0xE0, 0xD5)

def add_callout(doc, title, text, accent_color=CORAL):
    """Add an accented callout box using a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, 'FFF5F0')

    p = cell.paragraphs[0]
    run = p.add_run(title + '  ')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = accent_color
    run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = MEDIUM
    run.font.name = 'Calibri'

    # Left border accent
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:color="{str(accent_color)}" w:space="0"/>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    doc.add_paragraph()

def hex_color(rgb_color):
    return f'{rgb_color.red:02X}{rgb_color.green:02X}{rgb_color.blue:02X}'

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# BUILD THE DOCUMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default paragraph style
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = DARK

# ── COVER / TITLE SECTION ──────────────────────────────────────────

doc.add_paragraph()  # Top spacer

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FRIENDLY NATHAN')
run.font.size = Pt(36)
run.font.color.rgb = CORAL
run.font.name = 'Calibri'
run.bold = True
title.paragraph_format.space_after = Pt(2)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Your AI Workflow Buddy for n8n')
run.font.size = Pt(16)
run.font.color.rgb = TEAL
run.font.name = 'Calibri'
sub.paragraph_format.space_after = Pt(24)

# Version badge line
badge = doc.add_paragraph()
badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = badge.add_run('Release Notes  |  Version 1.0.0  |  February 2026')
run.font.size = Pt(11)
run.font.color.rgb = LIGHT_TEXT
run.font.name = 'Calibri'
badge.paragraph_format.space_after = Pt(6)

# Tagline
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tag.add_run('Describe what you want. Nathan builds it. One click to deploy.')
run.font.size = Pt(13)
run.font.color.rgb = MEDIUM
run.font.name = 'Calibri'
run.italic = True
tag.paragraph_format.space_after = Pt(30)

add_divider(doc)

# ── 1. WHAT IS FRIENDLY NATHAN ─────────────────────────────────────

add_styled_heading(doc, '1. What is Friendly Nathan?', level=1, space_before=24)

add_body(doc, (
    'Friendly Nathan is an AI-powered natural language workflow generator for n8n. '
    'Instead of manually dragging nodes, configuring parameters, and connecting integrations, '
    'you describe what you want in plain English and Nathan builds a complete, production-ready '
    'n8n workflow in seconds \u2014 then deploys it to your instance with a single click.'
))

add_callout(doc, 'Example:', (
    '"Every morning at 9am, check my Gmail for unread emails with attachments, '
    'save the attachments to Google Drive, and log a summary row in Google Sheets"'
), CORAL)

add_body(doc, (
    'Nathan turns that sentence into a fully configured n8n workflow with 6 nodes, '
    '5 connections, proper credential placeholders, trigger scheduling, and error handling.'
))

# ── 2. HOW IT WORKS ────────────────────────────────────────────────

add_styled_heading(doc, '2. How It Works', level=1)

add_body(doc, 'The generation pipeline follows a 6-stage process:', bold=True)

add_feature_table(doc,
    headers=['Stage', 'What Happens', 'Technology'],
    rows=[
        ['1. Connect', 'User provides n8n URL + API key. Nathan validates the connection.', 'REST API validation'],
        ['2. Discover', 'Nathan queries the n8n instance for all available nodes, caches the catalog.', 'MCP Server + 1hr cache'],
        ['3. Analyze', 'Gemini AI parses the natural language description, extracts intent, parameters, services.', 'Google Gemini AI'],
        ['4. Generate', 'AI builds the complete workflow JSON: nodes, connections, parameters, positions.', 'Gemini + rule engine'],
        ['5. Improve', 'Gap detector runs up to 3 refinement passes, auto-fixing missing triggers, outputs, configs.', 'Auto-improvement loop'],
        ['6. Deploy', 'One-click push to the n8n instance. Returns a direct URL to the live workflow.', 'n8n REST API'],
    ],
    col_widths=[2.2, 8.5, 4.0]
)

# ── 3. THE AI ENGINE ───────────────────────────────────────────────

add_styled_heading(doc, '3. The AI Engine', level=1)

add_styled_heading(doc, '3.1 Gemini Integration', level=2, color=TEAL)

add_body(doc, (
    'Nathan uses Google Gemini as its AI backbone with a smart fallback chain '
    'that ensures generation never fails due to model availability:'
))

add_feature_table(doc,
    headers=['Priority', 'Model', 'API Version', 'Use Case'],
    rows=[
        ['Primary', 'gemini-3-flash-preview', 'v1beta', 'Fastest, most capable for workflow generation'],
        ['Fallback 1', 'gemini-1.5-flash', 'v1', 'Fast, reliable fallback'],
        ['Fallback 2', 'gemini-1.5-pro', 'v1', 'Higher quality for complex workflows'],
        ['Fallback 3', 'gemini-1.0-pro', 'v1', 'Legacy fallback for maximum compatibility'],
    ],
    col_widths=[2.0, 4.5, 2.0, 6.5]
)

add_styled_heading(doc, '3.2 Three-Tier Prompt Architecture', level=2, color=TEAL)

add_body(doc, 'Every generation request is processed through a layered prompt system:')

add_bullet(doc, ' \u2014 Extracts sender emails, time ranges, channel IDs, spreadsheet IDs, and intent from natural language', bold_prefix='Tier 1: Intent Extraction')
add_bullet(doc, ' \u2014 Injects the 40 most relevant nodes from the user\'s n8n instance, with full parameter schemas, credential types, and default values', bold_prefix='Tier 2: Node Context')
add_bullet(doc, ' \u2014 Feeds patterns from past workflow fixes (up to 5 most relevant) to avoid repeating known issues', bold_prefix='Tier 3: Learning Guidance')

add_styled_heading(doc, '3.3 Chain+Model AI Pattern', level=2, color=TEAL)

add_body(doc, (
    'For any workflow involving AI processing, Nathan enforces a mandatory three-node pattern '
    'that follows n8n\'s LangChain architecture:'
))

add_feature_table(doc,
    headers=['Step', 'Node Type', 'n8n Package', 'Purpose'],
    rows=[
        ['1', 'Edit Fields', 'n8n-nodes-base.set', 'Prepares chatInput field from upstream data'],
        ['2', 'Basic LLM Chain', '@n8n/n8n-nodes-langchain.chainLlm', 'Orchestrates AI processing with prompt template'],
        ['3', 'Google Gemini Chat Model', '@n8n/n8n-nodes-langchain.lmChatGoogleGemini', 'Executes the AI inference'],
    ],
    col_widths=[1.2, 4.0, 6.5, 4.0]
)

add_body(doc, (
    'The chain connects via a special ai_model connection type (not the standard main connection), '
    'which is automatically configured by Nathan. This pattern prevents the most common mistakes '
    'users make when manually building AI workflows in n8n.'
))

# ── 4. AUTO-IMPROVEMENT ENGINE ─────────────────────────────────────

add_styled_heading(doc, '4. Auto-Improvement Engine', level=1)

add_body(doc, (
    'Nathan doesn\'t just generate workflows \u2014 it iteratively improves them. '
    'Two distinct improvement systems run after initial generation:'
))

add_styled_heading(doc, '4.1 AI Verification Loop (2 passes)', level=2, color=TEAL)

add_bullet(doc, ' \u2014 Gemini re-reads the generated workflow against the original description, identifies critical issues (missing nodes, wrong parameters, incomplete logic), and applies fixes automatically.', bold_prefix='Pass 1: Verify & Fix')
add_bullet(doc, ' \u2014 The fixed workflow is verified again. Any remaining suggestions are applied, and every issue/fix pair is recorded in the learning system for future improvements.', bold_prefix='Pass 2: Re-verify & Learn')

add_styled_heading(doc, '4.2 Gap Detection Loop (up to 3 passes)', level=2, color=TEAL)

add_body(doc, 'The gap detector scans for 7 specific problem categories:')

add_feature_table(doc,
    headers=['Gap Type', 'Severity', 'What It Detects', 'Auto-Fix'],
    rows=[
        ['Missing Trigger', 'High', 'Workflow has no entry point', 'Adds Schedule, Webhook, or Manual trigger'],
        ['Missing Output', 'High', 'Workflow processes data but sends it nowhere', 'Adds Gmail, Sheets, Slack, or notification node'],
        ['Vague Source', 'Medium', 'Data source lacks specifics (sender, subject, date range)', 'Prompts for details, adds filter nodes'],
        ['Missing AI Details', 'Medium', 'AI node without clear task definition', 'Suggests: Summarize, Extract, Classify, Generate'],
        ['Unclear Condition', 'Medium', 'If/Switch node without defined logic', 'Adds field comparison or time-based filtering'],
        ['Missing Transform', 'Low', 'Multiple items flow without aggregation', 'Adds Item Lists, Merge, or Set nodes'],
        ['Missing Aggregation', 'Low', 'Data needs combining before AI processing', 'Adds Edit Fields with aggregation expression'],
    ],
    col_widths=[3.0, 1.8, 5.5, 5.5]
)

add_callout(doc, 'Result:', (
    'Across both improvement systems, a single workflow generation can involve up to 5 total AI passes '
    '(2 verification + 3 gap detection), each building on the previous result. The learning system '
    'ensures that fixes discovered in one generation improve all future generations.'
), TEAL)

# ── 5. NODE DISCOVERY & CACHING ────────────────────────────────────

add_styled_heading(doc, '5. Intelligent Node Discovery', level=1)

add_body(doc, (
    'Nathan dynamically discovers what nodes are available on each user\'s n8n instance, '
    'ensuring generated workflows only reference nodes that actually exist:'
))

add_bullet(doc, ' \u2014 Primary discovery method using the Model Context Protocol for structured, fast node enumeration.', bold_prefix='MCP Server: ')
add_bullet(doc, ' \u2014 If MCP is unavailable, queries /api/v1/nodes directly with a 15-second timeout.', bold_prefix='Direct n8n API: ')
add_bullet(doc, ' \u2014 In-memory (session) + database (Prisma, 1-hour TTL per instance URL). Graceful degradation: serves expired cache if a fresh request fails.', bold_prefix='Multi-Layer Cache: ')

add_body(doc, (
    'A relevance scoring system (3 points for direct matches, 2 for short name matches, '
    '1 for multi-token matches) selects the top 40 most relevant nodes for each request, '
    'keeping the AI context focused and accurate.'
))

# ── 6. THE FRONTEND EXPERIENCE ─────────────────────────────────────

add_styled_heading(doc, '6. The Frontend Experience', level=1)

add_styled_heading(doc, '6.1 Meet Nathan \u2014 Your Workflow Buddy', level=2, color=TEAL)

add_body(doc, (
    'Nathan is represented by an expressive blob avatar with 6 distinct mood states, '
    'each with unique animations, colors, and expressions:'
))

add_feature_table(doc,
    headers=['Mood', 'When It Appears', 'Visual Details'],
    rows=[
        ['Idle', 'Default state, waiting for input', 'Gentle breathing animation (3s), warm orange glow, neutral smile'],
        ['Listening', 'User is typing in the text field', 'Focused narrowed eyes, alert posture, subtle lean forward'],
        ['Thinking', 'Workflow is being generated', 'Eyes look to the right, questioning mouth, 2s rotation animation'],
        ['Excited', 'Workflow preview is ready', 'Enlarged eyes (1.3x), big filled smile, 0.4s bounce, sparkle particles'],
        ['Success', 'Workflow deployed to n8n', 'Bright green glow (#34d399), sparkling circles, celebration expression'],
        ['Error', 'Something went wrong', 'Red/pink glow (#fb7185), downturned mouth, lowered eyes'],
    ],
    col_widths=[2.0, 5.0, 8.5]
)

add_styled_heading(doc, '6.2 Bright Sunny Theme', level=2, color=TEAL)

add_body(doc, 'The UI uses a distinctive warm cream palette designed to feel friendly and inviting:')

add_feature_table(doc,
    headers=['Element', 'Color', 'Usage'],
    rows=[
        ['Background', 'Sunny Cream (#FBF7F0)', 'Page background with floating dot pattern'],
        ['Primary', 'Coral Orange (#F06A3E)', 'Buttons, links, Nathan\'s glow, focus rings'],
        ['Secondary', 'Bright Teal (#14B8A6)', 'Success states, accents, "workflow buddy" underline'],
        ['Accent', 'Sunny Yellow (#E8AA0A)', 'Highlights, hover states, sparkle particles'],
        ['Typography', 'Bricolage Grotesque / DM Sans / JetBrains Mono', 'Display / body / code fonts via Google Fonts'],
    ],
    col_widths=[2.5, 5.5, 7.5]
)

add_styled_heading(doc, '6.3 Real-Time Progress', level=2, color=TEAL)

add_body(doc, (
    'Workflow generation streams progress updates via WebSocket (Socket.io). '
    'Users see a live progress bar with descriptive messages at each stage:'
))

add_bullet(doc, '10% \u2014 "Starting workflow generation..."')
add_bullet(doc, '15% \u2014 "Discovering available nodes..."')
add_bullet(doc, '30% \u2014 "Analyzing description..."')
add_bullet(doc, '45% \u2014 "Using AI to understand your request..."')
add_bullet(doc, '55% \u2014 "Validating workflow..."')
add_bullet(doc, '60% \u2014 "Creating workflow in n8n..."')
add_bullet(doc, '100% \u2014 "Workflow created successfully!"')

add_body(doc, 'Estimated time remaining is calculated dynamically based on current progress.', color=LIGHT_TEXT, size=9.5)

add_styled_heading(doc, '6.4 Smart Suggestion Chips', level=2, color=TEAL)

add_body(doc, 'The landing page offers emoji-enhanced suggestion chips to inspire first-time users:')

add_feature_table(doc,
    headers=['Emoji', 'Suggestion'],
    rows=[
        ['\U0001F4E7', 'Read new Gmail emails and save to Google Sheets'],
        ['\U0001F916', 'Summarize Google Docs with AI daily'],
        ['\U0001F517', 'Webhook receives data, process with Gemini, store in database'],
        ['\U0001F4CA', 'Monitor Google Sheet changes and send email alerts'],
    ],
    col_widths=[1.5, 14.0]
)

# ── 7. TEMPLATE SYSTEM ─────────────────────────────────────────────

add_styled_heading(doc, '7. Template System', level=1)

add_body(doc, (
    '10 pre-built workflow templates let users skip the description entirely. '
    'Each template has fillable fields with validation, difficulty ratings, and estimated setup time:'
))

add_feature_table(doc,
    headers=['#', 'Template', 'Category', 'Difficulty', 'Time'],
    rows=[
        ['1', 'Email Summary to Slack', 'Email Automation', 'Beginner', '2 min'],
        ['2', 'Emails to Spreadsheet', 'Email Automation', 'Beginner', '2 min'],
        ['3', 'Webhook to Slack Alert', 'Data Collection', 'Beginner', '1 min'],
        ['4', 'Form Submissions to Sheet', 'Data Collection', 'Beginner', '2 min'],
        ['5', 'Daily Report via Email', 'Reporting', 'Intermediate', '3 min'],
        ['6', 'API Monitor & Alert', 'Monitoring', 'Intermediate', '3 min'],
        ['7', 'Content Analyzer', 'AI Processing', 'Intermediate', '3 min'],
        ['8', 'Spreadsheet to DB Sync', 'Data Sync', 'Advanced', '5 min'],
        ['9', 'Daily Slack Reminder', 'Notifications', 'Beginner', '1 min'],
        ['10', 'Automated Email Response', 'Email Automation', 'Advanced', '4 min'],
    ],
    col_widths=[0.8, 4.5, 3.0, 2.5, 1.5]
)

add_body(doc, (
    'Templates support 7 field types: text, email, URL, Slack channel, number, select (dropdown), and textarea. '
    'Each field includes validation rules, placeholder text, and sensible defaults.'
))

# ── 8. MCP SERVER ──────────────────────────────────────────────────

add_styled_heading(doc, '8. MCP Server (Model Context Protocol)', level=1)

add_body(doc, (
    'Nathan includes a standalone MCP server that exposes n8n workflow capabilities '
    'to any MCP-compatible AI system (Claude, GPT, etc.). This enables external AI agents '
    'to discover, create, manage, and execute n8n workflows programmatically.'
))

add_feature_table(doc,
    headers=['Tool', 'Description'],
    rows=[
        ['n8n_list_nodes', 'List all available nodes with category/search filters'],
        ['n8n_get_node_details', 'Get detailed configuration for a specific node type'],
        ['n8n_suggest_workflow', 'Analyze a description and suggest appropriate nodes'],
        ['n8n_create_workflow', 'Create a new workflow on the n8n instance'],
        ['n8n_get_workflow', 'Retrieve a workflow\'s full configuration'],
        ['n8n_list_workflows', 'List workflows with active/inactive filters'],
        ['n8n_activate_workflow', 'Toggle a workflow\'s active state'],
        ['n8n_delete_workflow', 'Delete a workflow from the instance'],
        ['n8n_execute_workflow', 'Trigger manual execution of a workflow'],
        ['n8n_get_node_types', 'Get detailed node schemas with all properties (cached 1hr)'],
        ['n8n_refresh_node_cache', 'Force refresh the node discovery cache'],
        ['n8n_get_cache_status', 'Check cache age and remaining TTL'],
    ],
    col_widths=[4.0, 11.5]
)

# ── 9. SECURITY ────────────────────────────────────────────────────

add_styled_heading(doc, '9. Security & Reliability', level=1)

add_feature_table(doc,
    headers=['Layer', 'Implementation', 'Details'],
    rows=[
        ['Encryption', 'AES-256-CBC', 'Random IV per operation. Stored credentials encrypted at rest.'],
        ['Authentication', 'JWT + bcrypt', 'Token-based auth with salted password hashing (10 rounds).'],
        ['Input Validation', 'Zod schemas', 'Every endpoint validates input. Max lengths, format checks, type safety.'],
        ['HTTP Security', 'Helmet.js', 'Security headers: CSP, HSTS, X-Frame-Options, referrer policy.'],
        ['Rate Limiting', 'express-rate-limit', '50 req/15min general, 10 req/15min for generation endpoints.'],
        ['Access Control', 'RBAC', 'Employee and Admin roles with middleware-enforced route protection.'],
        ['Startup Validation', 'Fail-fast', 'Server refuses to start without JWT_SECRET and ENCRYPTION_KEY.'],
        ['Error Boundary', 'React ErrorBoundary', 'Frontend catches crashes gracefully with recovery UI.'],
        ['XSS Prevention', 'Safe rendering', 'JSON syntax highlighting uses React elements, not innerHTML.'],
    ],
    col_widths=[2.5, 3.5, 9.5]
)

# ── 10. TECH STACK ─────────────────────────────────────────────────

add_styled_heading(doc, '10. Technology Stack', level=1)

add_feature_table(doc,
    headers=['Layer', 'Technologies', 'Versions'],
    rows=[
        ['Frontend', 'React, TypeScript, TailwindCSS, Vite, Motion, Socket.io Client', '18.2, 5.3, 3.4, 5.0, 12.33, 4.6'],
        ['Backend', 'Express, TypeScript, Prisma ORM, Socket.io, Zod', '4.18, 5.3, 5.8, 4.6, 3.22'],
        ['AI Engine', 'Google Gemini (4-model fallback chain)', 'gemini-3-flash-preview (primary)'],
        ['Database', 'PostgreSQL + Prisma migrations', '15+'],
        ['MCP Server', 'Model Context Protocol SDK, Axios', '1.25, 1.6'],
        ['Security', 'Helmet, JWT, AES-256-CBC, bcrypt, express-rate-limit', '7.1, 9.0, -, 5.1, 7.1'],
        ['DevOps', 'Docker, Docker Compose, Nginx, GitHub', 'Latest'],
        ['Typography', 'Bricolage Grotesque, DM Sans, JetBrains Mono', 'Google Fonts'],
    ],
    col_widths=[2.5, 8.5, 5.0]
)

# ── 11. ARCHITECTURE ───────────────────────────────────────────────

add_styled_heading(doc, '11. System Architecture', level=1)

add_body(doc, 'Friendly Nathan is a monorepo with three independent services:', bold=True)

add_feature_table(doc,
    headers=['Service', 'Port', 'Responsibility', 'Key Modules'],
    rows=[
        ['Frontend', ':5173', 'User interface, real-time progress, workflow preview', '16 pages, 11 components, auth context'],
        ['Backend', ':3000', 'AI pipeline, workflow generation, n8n deployment', '8 controllers, 11 services, 8 routes'],
        ['MCP Server', 'stdio', 'External AI integration, node discovery', '12 tools, node caching, suggestion engine'],
    ],
    col_widths=[2.5, 1.5, 6.0, 5.5]
)

add_body(doc, 'Communication flow:', bold=True, space_after=3)
add_bullet(doc, 'Frontend \u2194 Backend: HTTP REST + WebSocket (Socket.io)')
add_bullet(doc, 'Backend \u2194 n8n: REST API with encrypted credentials')
add_bullet(doc, 'Backend \u2194 MCP: Model Context Protocol over stdio')
add_bullet(doc, 'Backend \u2194 Gemini: REST API (v1/v1beta) with model fallback chain')
add_bullet(doc, 'Backend \u2194 PostgreSQL: Prisma ORM with connection pooling')

# ── 12. KEY METRICS ────────────────────────────────────────────────

add_styled_heading(doc, '12. Key Metrics', level=1)

add_feature_table(doc,
    headers=['Metric', 'Value'],
    rows=[
        ['AI model fallback chain depth', '4 models'],
        ['Auto-improvement max iterations', '5 (2 AI verification + 3 gap detection)'],
        ['Common nodes pool', '23 pre-configured essential nodes'],
        ['Node cache TTL', '60 minutes per instance'],
        ['Preview cache capacity', '100 concurrent (15-min TTL)'],
        ['Gap detection categories', '7 types with auto-fix'],
        ['Workflow templates', '10 across 7 categories'],
        ['Template field types', '7 (text, email, URL, channel, number, select, textarea)'],
        ['MCP tools exposed', '12 for external AI integration'],
        ['Security layers', '9 (encryption, auth, validation, headers, rate limiting, RBAC, startup, error boundary, XSS)'],
        ['Frontend pages', '16 route components'],
        ['Backend services', '11 business logic modules'],
        ['Supported fonts', '3 families (display, body, monospace)'],
    ],
    col_widths=[6.0, 9.5]
)

# ── 13. WHAT MAKES NATHAN UNIQUE ───────────────────────────────────

add_styled_heading(doc, '13. What Makes Nathan Unique', level=1)

unique_features = [
    ('Cascading AI Verification', 'Leverages Gemini for both generation AND verification with a learning feedback loop.'),
    ('Adaptive Node Discovery', 'MCP + Direct API + Database caching with intelligent fallback \u2014 never fails to discover.'),
    ('Self-Improving System', 'Every workflow fix is recorded and used to improve future generations automatically.'),
    ('Chain+Model Standardization', 'Mandatory AI pattern prevents the #1 mistake users make with n8n AI nodes.'),
    ('Expressive Avatar System', '6 mood states with smooth CSS animations, sparkle particles, and color-coded feedback.'),
    ('Smart Google Substitution', 'Requests for Outlook, Excel, Airtable, Notion are automatically mapped to Google equivalents.'),
    ('Zero-Config Templates', '10 templates from Beginner to Advanced with fillable fields, validation, and one-click deploy.'),
    ('Real-Time Generation', 'WebSocket progress streaming with estimated time \u2014 users never stare at a loading spinner.'),
    ('Graceful Degradation', 'Works without Gemini (rule-based fallback), survives node discovery failures, falls back across 4 AI models.'),
    ('MCP Integration', 'The only n8n workflow generator that exposes capabilities via Model Context Protocol for external AI agents.'),
]

for i, (title, desc) in enumerate(unique_features, 1):
    add_bullet(doc, f'  {desc}', bold_prefix=f'{i}. {title}')

# ── FOOTER ─────────────────────────────────────────────────────────

add_divider(doc)

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Built by the RISE Team')
run.font.size = Pt(11)
run.font.color.rgb = CORAL
run.font.name = 'Calibri'
run.bold = True

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('Making workflow automation accessible to everyone')
run.font.size = Pt(10)
run.font.color.rgb = LIGHT_TEXT
run.font.name = 'Calibri'
run.italic = True

# ── SAVE ───────────────────────────────────────────────────────────

output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'Friendly-Nathan-v1.0-Release-Notes.docx')
doc.save(output_path)
print(f'Release notes saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
